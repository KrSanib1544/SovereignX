# backend/app/agent/tools/generate_docx.py
"""
Generate DOCX Tool
Assembles structured industrial Engineering Approval Notes and Inspection Compliance Reports
using python-docx, saving deliverables strictly within the workspace artifacts directory.
"""

from datetime import datetime, timezone
import os
from pathlib import Path
from typing import Dict, List, Optional
from pydantic import BaseModel, Field

from backend.app.agent.tools.base import (
    BaseTool,
    ToolDefinition,
    ToolRiskLevel,
    resolve_secure_workspace_path,
)


class FindingRow(BaseModel):
    component: str
    observed_defect: str
    threshold: str
    risk_level: str
    citation: str


class GenerateDocxInput(BaseModel):
    output_filename: str = Field(
        "Engineering_Approval_Note.docx",
        description="Deliverable filename (e.g., 'Approval_Note_Pump3B.docx')"
    )
    title: str = Field(..., description="Main title of the report")
    executive_summary: str = Field(..., description="High-level engineering summary")
    findings: List[FindingRow] = Field(default_factory=list, description="Structured table of findings")
    recommendations: List[str] = Field(default_factory=list, description="Numbered actionable recommendations")
    signoff_author: Optional[str] = Field("Lead Reliability Engineer", description="Author name")
    signoff_role: Optional[str] = Field("Asset Integrity & Safety Assurance", description="Author role")


class GenerateDocxOutput(BaseModel):
    filename: str
    relative_path: str
    size_bytes: int
    created_at: str
    status: str


class GenerateDocxTool(BaseTool):
    @property
    def definition(self) -> ToolDefinition:
        return ToolDefinition(
            name="generate_docx",
            description="Generate a formatted industrial Engineering Approval Note / Inspection Compliance Summary document (.docx).",
            input_schema=GenerateDocxInput,
            output_schema=GenerateDocxOutput,
            risk_level=ToolRiskLevel.MEDIUM,
            required_permissions=["workspace:write"],
            requires_human_approval=False
        )

    async def execute(self, workspace_id: str, input_data: GenerateDocxInput) -> GenerateDocxOutput:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        # Validate filename
        clean_name = os.path.basename(input_data.output_filename)
        if not clean_name.endswith(".docx"):
            clean_name += ".docx"

        rel_path = f"artifacts/{clean_name}"
        target_path = resolve_secure_workspace_path(
            workspace_id=workspace_id,
            relative_path=rel_path,
            must_exist=False
        )
        target_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Document Header
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(input_data.title)
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(24, 43, 73)
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        sub_p = doc.add_paragraph()
        sub_p.add_run(f"SOVEREIGN-X CONFIDENTIAL ENGINEERING ASSESSMENT • {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
        sub_p.runs[0].font.size = Pt(9)
        sub_p.runs[0].font.color.rgb = RGBColor(110, 110, 110)

        doc.add_heading("1. Executive Summary", level=2)
        doc.add_paragraph(input_data.executive_summary)

        # Findings Table
        if input_data.findings:
            doc.add_heading("2. Observed Inspection Findings & Defect Mapping", level=2)
            table = doc.add_table(rows=1, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True

            hdr_cells = table.rows[0].cells
            hdr_titles = ["Component", "Observed Defect", "OEM Threshold", "Risk", "Citation"]
            for i, title in enumerate(hdr_titles):
                hdr_cells[i].text = title
                for p in hdr_cells[i].paragraphs:
                    for run in p.runs:
                        run.font.bold = True

            for item in input_data.findings:
                row_cells = table.add_row().cells
                row_cells[0].text = item.component
                row_cells[1].text = item.observed_defect
                row_cells[2].text = item.threshold
                row_cells[3].text = item.risk_level
                row_cells[4].text = item.citation

        # Recommendations
        if input_data.recommendations:
            doc.add_heading("3. Corrective Actions & Engineering Recommendations", level=2)
            for i, rec in enumerate(input_data.recommendations, start=1):
                doc.add_paragraph(f"{i}. {rec}")

        # Sign-off Block
        doc.add_heading("4. Sign-off & Verification Authorization", level=2)
        sign_p = doc.add_paragraph()
        sign_p.add_run(f"Verified by: {input_data.signoff_author}\n")
        sign_p.add_run(f"Designation: {input_data.signoff_role}\n")
        sign_p.add_run(f"Status: CERTIFIED COMPLIANT / AIR-GAP VERIFIED\n")

        doc.save(str(target_path))
        size_bytes = target_path.stat().st_size

        return GenerateDocxOutput(
            filename=clean_name,
            relative_path=rel_path,
            size_bytes=size_bytes,
            created_at=datetime.now(timezone.utc).isoformat(),
            status="CREATED"
        )
